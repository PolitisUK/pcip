"""Extract an approved Word policy into canonical repository Markdown.

This is an authoring-time tool only. The deployed application reads the
generated Markdown; it never converts DOCX files at runtime.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def _normalise(value: str) -> str:
    return " ".join(value.split())


def extract(source: Path, destination: Path, source_label: str) -> None:
    document = Document(source)
    lines = [f"<!-- Canonical source: {source_label} -->", ""]
    paragraphs = {id(paragraph._p): paragraph for paragraph in document.paragraphs}
    tables = {id(table._tbl): table for table in document.tables}
    table_count = 0
    for child in document.element.body.iterchildren():
        paragraph = paragraphs.get(id(child))
        if paragraph is not None:
            text = _normalise(paragraph.text)
            if not text:
                continue
            if paragraph.style.name == "Heading 1":
                lines.extend((f"## {text}", ""))
            elif paragraph.style.name == "Heading 2":
                lines.extend((f"### {text}", ""))
            elif paragraph.style.name.startswith("List Bullet"):
                lines.extend((f"- {text}", ""))
            else:
                lines.extend((text, ""))
            continue
        table = tables.get(id(child))
        if table is None:
            continue
        rows = [" — ".join(_normalise(cell.text) for cell in row.cells) for row in table.rows]
        if rows:
            heading = "Publication details" if table_count == 0 else "Important information"
            lines.extend((f"## {heading}", ""))
            for row in rows:
                lines.extend((f"> {row}", ""))
        table_count += 1
    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("destination", type=Path)
    parser.add_argument("--source-label", required=True)
    args = parser.parse_args()
    extract(args.source, args.destination, args.source_label)


if __name__ == "__main__":
    main()
